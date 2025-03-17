import os
import datetime
import pandas as pd
import xmlrpc.client
import io
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from dotenv import load_dotenv
import streamlit as st
from collections import defaultdict

# Load environment variables from .env file
load_dotenv()

# Odoo credentials
ODOO_URL = st.secrets["odoo"]["ODOO_URL"]
ODOO_DB = st.secrets["odoo"]["ODOO_DB"]
ODOO_USERNAME = st.secrets["odoo"]["ODOO_USERNAME"]
ODOO_PASSWORD = st.secrets["odoo"]["ODOO_PASSWORD"]

def set_collapsible(paragraph):
    """
    Attempts to add a collapsible property to a heading by injecting <w:collapse>.
    For this to work, the heading must use a built-in style (like 'Heading 1')
    and be viewed in a Word version that supports collapsible headings.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    collapse = OxmlElement('w:collapse')
    collapse.set(qn('w:val'), "true")
    pPr.append(collapse)

def set_cell_shading(cell, fill="D9E1F2"):
    """
    Sets the background shading of a cell.
    Default fill color is a light blue.
    """
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margin(cell, margin=100):
    """
    Sets uniform cell margins (in dxa units; 100 dxa ≃ 0.07 inches) for the given cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:w="{margin}" w:type="dxa"/>'
        f'<w:left w:w="{margin}" w:type="dxa"/>'
        f'<w:bottom w:w="{margin}" w:type="dxa"/>'
        f'<w:right w:w="{margin}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)

def set_column_widths(table, widths):
    """
    Sets the width for each column in the table. 
    widths should be a list of Inches objects matching the number of columns.
    """
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            set_cell_margin(cell, margin=100)

def authenticate_odoo():
    """Authenticate with Odoo and return UID, models object."""
    common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
    uid = common.authenticate(ODOO_DB, ODOO_USERNAME, ODOO_PASSWORD, {})
    models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")
    return uid, models

def get_sunday_friday_range():
    """Returns (start_of_week, end_of_week) for the current work week (Sunday–Friday)."""
    today = datetime.date.today()
    diff = today.weekday() - 6
    if diff < 0:
        diff += 7
    start_of_week = today - datetime.timedelta(days=diff)
    end_of_week = start_of_week + datetime.timedelta(days=5)
    return start_of_week, end_of_week

def get_designer_ids_from_planning(models, uid, start_date, end_date):
    """Queries planning.slot for the given date range and returns IDs of designers."""
    slots = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'planning.slot', 'search_read',
        [[('start_datetime', '>=', start_date), ('end_datetime', '<=', end_date)]],
        {'fields': ['resource_id']}
    )
    resource_ids = {
        slot['resource_id'][0] if isinstance(slot.get('resource_id'), list) else slot.get('resource_id')
        for slot in slots if slot.get('resource_id')
    }
    if not resource_ids:
        return []
    employees = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'hr.employee', 'search_read',
        [[('id', 'in', list(resource_ids))]],
        {'fields': ['id', 'name', 'job_title']}
    )
    designer_ids = [emp['id'] for emp in employees if 'designer' in (emp.get('job_title') or '').lower()]
    return designer_ids

def read_employee_info(models, uid, employee_ids):
    """Retrieves full employee records for the given IDs."""
    if not employee_ids:
        return []
    return models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'hr.employee', 'search_read',
        [[('id', 'in', employee_ids)]],
        {'fields': ['id', 'name', 'job_title', 'user_id']}
    )

def get_all_timesheet_hours(models, uid, designer_ids, start_date, end_date):
    """Retrieves timesheet hours for the given designer IDs."""
    if not designer_ids:
        return {}
    timesheets = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'account.analytic.line', 'search_read',
        [[('employee_id', 'in', designer_ids),
          ('date', '>=', start_date),
          ('date', '<=', end_date)]],
        {'fields': ['employee_id', 'unit_amount']}
    )
    timesheet_dict = defaultdict(float)
    for ts in timesheets:
        emp_field = ts.get('employee_id')
        if emp_field:
            emp_id = emp_field[0] if isinstance(emp_field, list) else emp_field
            timesheet_dict[emp_id] += float(ts.get('unit_amount', 0))
    return dict(timesheet_dict)

def get_all_scheduled_data(models, uid, designer_ids, start_date, end_date):
    """Retrieves scheduling data (hours and projects) from planning.slot."""
    if not designer_ids:
        return {}
    slots = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'planning.slot', 'search_read',
        [[('resource_id', 'in', designer_ids),
          ('start_datetime', '>=', start_date),
          ('end_datetime', '<=', end_date)]],
        {'fields': ['resource_id', 'start_datetime', 'end_datetime', 'project_id']}
    )
    scheduled_data = {}
    for slot in slots:
        res_field = slot.get('resource_id')
        if not res_field:
            continue
        emp_id = res_field[0] if isinstance(res_field, list) else res_field
        if emp_id not in scheduled_data:
            scheduled_data[emp_id] = {'hours': 0.0, 'projects': set()}
        start = pd.to_datetime(slot['start_datetime'])
        end = pd.to_datetime(slot['end_datetime'])
        hours = (end - start).total_seconds() / 3600.0
        scheduled_data[emp_id]['hours'] += hours
        project_field = slot.get('project_id')
        if project_field:
            project_name = project_field[1] if isinstance(project_field, list) else str(project_field)
            scheduled_data[emp_id]['projects'].add(project_name)
    return scheduled_data

def get_subtask_service_categories(models, uid, designer_ids, start_date, end_date):
    """For planning slots with subtask references, retrieves the service category."""
    if not designer_ids:
        return {}
    main_slots = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'planning.slot', 'search_read',
        [[('resource_id', 'in', designer_ids),
          ('start_datetime', '>=', start_date),
          ('end_datetime', '<=', end_date),
          ('x_studio_sub_task_1', '!=', False)]],
        {'fields': ['resource_id', 'x_studio_sub_task_1']}
    )
    emp_task_pairs = []
    for slot in main_slots:
        res_field = slot.get('resource_id')
        if not res_field:
            continue
        emp_id = res_field[0] if isinstance(res_field, list) else res_field
        subtask_field = slot.get('x_studio_sub_task_1')
        if not subtask_field:
            continue
        task_id = subtask_field[0] if isinstance(subtask_field, list) else subtask_field
        emp_task_pairs.append((emp_id, task_id))
    unique_task_ids = list({task_id for (_, task_id) in emp_task_pairs})
    if not unique_task_ids:
        return {}
    tasks_data = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'project.task', 'read',
        [unique_task_ids],
        {'fields': ['x_studio_service_category_1']}
    )
    task_cat_map = {}
    for task in tasks_data:
        cat_field = task.get('x_studio_service_category_1')
        if cat_field:
            cat_name = cat_field[1] if isinstance(cat_field, list) else str(cat_field)
            task_cat_map[task['id']] = cat_name
    categories_dict = {}
    for emp_id, task_id in emp_task_pairs:
        cat_name = task_cat_map.get(task_id)
        if cat_name:
            if emp_id not in categories_dict:
                categories_dict[emp_id] = set()
            categories_dict[emp_id].add(cat_name)
    return categories_dict

def get_parent_task_due_dates(models, uid, designer_ids, start_date, end_date):
    """For planning slots with parent task references, retrieves and converts the deadline."""
    if not designer_ids:
        return {}
    main_slots = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'planning.slot', 'search_read',
        [[('resource_id', 'in', designer_ids),
          ('start_datetime', '>=', start_date),
          ('end_datetime', '<=', end_date),
          ('x_studio_parent_task', '!=', False)]],
        {'fields': ['resource_id', 'x_studio_parent_task']}
    )
    emp_task_pairs = []
    for slot in main_slots:
        res_field = slot.get('resource_id')
        if not res_field:
            continue
        emp_id = res_field[0] if isinstance(res_field, list) else res_field
        parent_field = slot.get('x_studio_parent_task')
        if not parent_field:
            continue
        task_id = parent_field[0] if isinstance(parent_field, list) else parent_field
        emp_task_pairs.append((emp_id, task_id))
    unique_task_ids = list({task_id for (_, task_id) in emp_task_pairs})
    if not unique_task_ids:
        return {}
    tasks_data = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'project.task', 'read',
        [unique_task_ids],
        {'fields': ['x_studio_internal_due_date_1']}
    )
    task_due_map = {}
    for task in tasks_data:
        raw_date = task.get('x_studio_internal_due_date_1')
        if raw_date:
            try:
                dt_parsed = pd.to_datetime(raw_date)
                if dt_parsed.tzinfo is None:
                    dt_parsed = dt_parsed.replace(tzinfo=datetime.timezone.utc)
                dt_local = dt_parsed.astimezone(datetime.timezone(datetime.timedelta(hours=3)))
                due_date_str = dt_local.strftime('%Y-%m-%d %H:%M:%S')
            except Exception:
                due_date_str = str(raw_date)
            task_due_map[task['id']] = due_date_str
    deadlines_dict = {}
    for emp_id, task_id in emp_task_pairs:
        due_date_str = task_due_map.get(task_id)
        if due_date_str:
            if emp_id not in deadlines_dict:
                deadlines_dict[emp_id] = set()
            deadlines_dict[emp_id].add(due_date_str)
    return deadlines_dict

def get_project_breakdown(models, uid, designer_ids, start_date, end_date):
    """Builds a breakdown for each employee: {emp_id: {project_name: {project_type: count, ...}}}."""
    slots = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'planning.slot', 'search_read',
        [[('resource_id', 'in', designer_ids),
          ('start_datetime', '>=', start_date),
          ('end_datetime', '<=', end_date)]],
        {'fields': ['resource_id', 'project_id', 'x_studio_sub_task_1']}
    )
    task_ids = set()
    for slot in slots:
        if slot.get('x_studio_sub_task_1'):
            subtask_field = slot.get('x_studio_sub_task_1')
            task_id = subtask_field[0] if isinstance(subtask_field, list) else subtask_field
            task_ids.add(task_id)
    task_cat_map = {}
    if task_ids:
        tasks_data = models.execute_kw(
            ODOO_DB, uid, ODOO_PASSWORD,
            'project.task', 'read',
            [list(task_ids)],
            {'fields': ['x_studio_service_category_1']}
        )
        for task in tasks_data:
            cat_field = task.get('x_studio_service_category_1')
            if cat_field:
                cat_name = cat_field[1] if isinstance(cat_field, list) else str(cat_field)
            else:
                cat_name = None
            task_cat_map[task['id']] = cat_name
    breakdown = {}
    for slot in slots:
        res_field = slot.get('resource_id')
        if not res_field:
            continue
        emp_id = res_field[0] if isinstance(res_field, list) else res_field
        if emp_id not in breakdown:
            breakdown[emp_id] = {}
        project_field = slot.get('project_id')
        if not project_field:
            continue
        project_name = project_field[1] if isinstance(project_field, list) else str(project_field)
        project_type = None
        if slot.get('x_studio_sub_task_1'):
            subtask_field = slot.get('x_studio_sub_task_1')
            task_id = subtask_field[0] if isinstance(subtask_field, list) else subtask_field
            project_type = task_cat_map.get(task_id)
        type_key = project_type if project_type is not None else "No Type"
        if project_name not in breakdown[emp_id]:
            breakdown[emp_id][project_name] = {}
        breakdown[emp_id][project_name][type_key] = breakdown[emp_id][project_name].get(type_key, 0) + 1
    return breakdown

def format_project_breakdown_for_employee(breakdown_for_employee):
    """Formats a breakdown dictionary into a total count and a summary string."""
    total_count = sum(sum(types.values()) for types in breakdown_for_employee.values())
    project_strings = []
    for project_name, types in breakdown_for_employee.items():
        if len(types) == 1 and list(types.values())[0] == 1 and "No Type" in types:
            project_strings.append(project_name)
        else:
            details = ", ".join(f"{count} {type_name}" for type_name, count in types.items())
            project_strings.append(f"{project_name} ({details})")
    if len(project_strings) > 1:
        final_projects_str = ", ".join(project_strings[:-1]) + " and " + project_strings[-1]
    else:
        final_projects_str = project_strings[0] if project_strings else ""
    return total_count, final_projects_str

def aggregate_project_breakdowns(project_breakdown_dict):
    """
    Aggregates the project breakdown across all designers, grouping results by Project Type.
    
    Returns a dictionary where each key is a project type (or "No Type" if missing) and each value is a 
    dictionary mapping project names to their aggregated counts.
    """
    aggregated = {}
    for emp_id, emp_breakdown in project_breakdown_dict.items():
        for project_name, type_dict in emp_breakdown.items():
            for project_type, count in type_dict.items():
                key = project_type if project_type is not None else "No Type"
                if key not in aggregated:
                    aggregated[key] = {}
                aggregated[key][project_name] = aggregated[key].get(project_name, 0) + count
    return aggregated

def get_availability_guess_coded(designer_name, timesheet_hours, scheduled_hours):
    """Computes available hours and returns a guess string."""
    weekly_hours = 40
    available_hours = weekly_hours - (timesheet_hours + scheduled_hours)
    if available_hours < 0:
        available_hours = 0
    if available_hours >= 15:
        guess = "Fully Available"
    elif available_hours > 0:
        guess = "Partially Available"
    else:
        guess = "Not Available"
    return available_hours, guess

def create_deadline_pie_chart(deadline_list):
    """
    Creates a pie chart image from a list of deadline strings, with a larger figure size.
    - Next week (0-6 days) => red
    - Next 2 weeks (7-13 days) => yellow
    - Beyond 2 weeks (>=14 days) => green
    If no valid deadlines are found, creates a placeholder chart.
    Returns an image buffer.
    """
    now = datetime.datetime.now()
    red = yellow = green = 0
    for d_str in deadline_list:
        try:
            d = pd.to_datetime(d_str).to_pydatetime()
        except Exception:
            continue
        delta = (d - now).days
        if delta < 0:
            continue  # ignore past deadlines
        elif delta < 7:
            red += 1
        elif delta < 14:
            yellow += 1
        else:
            green += 1
    counts = [red, yellow, green]
    fig, ax = plt.subplots(figsize=(6, 6))
    if sum(counts) == 0:
        ax.pie([1], labels=["No deadlines"], colors=["gray"], autopct='%1.1f%%')
        ax.axis('equal')
    else:
        labels = [f"Next week ({red})", f"Next 2 weeks ({yellow})", f"Beyond 2 weeks ({green})"]
        colors = ["red", "yellow", "green"]
        ax.pie(counts, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
        ax.axis('equal')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=100)
    plt.close(fig)
    buf.seek(0)
    return buf

def get_deadlines_for_week(models, uid, designer_ids, start_date, end_date):
    """
    Retrieves planning.slot entries with a parent task (deadline) and corresponding project info.
    Returns a list of dicts with keys: 'designer', 'project', 'project_type', 'deadline'
    for deadlines within the next 7 days.
    """
    slots = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'planning.slot', 'search_read',
        [[('resource_id', 'in', designer_ids),
          ('start_datetime', '>=', start_date),
          ('end_datetime', '<=', end_date),
          ('x_studio_parent_task', '!=', False)]],
        {'fields': ['resource_id', 'project_id', 'x_studio_parent_task']}
    )
    task_ids = []
    for slot in slots:
        parent_task = slot.get('x_studio_parent_task')
        if parent_task:
            parent_task = parent_task[0] if isinstance(parent_task, list) else parent_task
            task_ids.append(parent_task)
    unique_task_ids = list(set(task_ids))
    if not unique_task_ids:
        return []
    tasks_data = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'project.task', 'read',
        [unique_task_ids],
        {'fields': ['x_studio_internal_due_date_1', 'x_studio_service_category_1']}
    )
    task_info = {}
    for task in tasks_data:
        raw_date = task.get('x_studio_internal_due_date_1')
        try:
            dt_parsed = pd.to_datetime(raw_date)
            if dt_parsed.tzinfo is None:
                dt_parsed = dt_parsed.replace(tzinfo=datetime.timezone.utc)
            dt_local = dt_parsed.astimezone(datetime.timezone(datetime.timedelta(hours=3)))
            deadline_str = dt_local.strftime('%Y-%m-%d %H:%M:%S')
        except Exception:
            deadline_str = str(raw_date)
        project_type = None
        cat_field = task.get('x_studio_service_category_1')
        if cat_field:
            project_type = cat_field[1] if isinstance(cat_field, list) else str(cat_field)
        task_info[task['id']] = {'deadline': deadline_str, 'project_type': project_type}
    employees = models.execute_kw(
        ODOO_DB, uid, ODOO_PASSWORD,
        'hr.employee', 'search_read',
        [[('id', 'in', designer_ids)]],
        {'fields': ['id', 'name']}
    )
    emp_names = {emp['id']: emp['name'] for emp in employees}
    deadlines_for_week = []
    now = datetime.datetime.now(datetime.timezone.utc)
    for slot in slots:
        designer_id = None
        if slot.get('resource_id'):
            designer_id = slot.get('resource_id')[0] if isinstance(slot.get('resource_id'), list) else slot.get('resource_id')
        parent_task = slot.get('x_studio_parent_task')
        if parent_task:
            parent_task = parent_task[0] if isinstance(parent_task, list) else parent_task
        if parent_task not in task_info:
            continue
        deadline_str = task_info[parent_task]['deadline']
        try:
            d_dt = pd.to_datetime(deadline_str)
        except Exception:
            continue
        delta_days = (d_dt - pd.Timestamp.now()).days
        if 0 <= delta_days < 7:
            record = {
                'designer': emp_names.get(designer_id, "Unknown"),
                'project': None,
                'project_type': task_info[parent_task]['project_type'],
                'deadline': deadline_str
            }
            project_field = slot.get('project_id')
            if project_field:
                record['project'] = project_field[1] if isinstance(project_field, list) else str(project_field)
            deadlines_for_week.append(record)
    return deadlines_for_week

def generate_better_word_doc(designer_info_list, aggregated_breakdown, aggregated_deadlines, deadlines_details):
    document = Document()
    # Main title as Heading 1 without the date range.
    heading_paragraph = document.add_heading("Designer Capacity and Availability", 1)
    set_collapsible(heading_paragraph)
    document.styles['Heading 1'].font.size = Pt(18)
    document.styles['Heading 1'].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    # Set base style.
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ----------------- Create a combined summary table at the top -----------------
    total_designers = len(designer_info_list)
    total_available_hours = sum(info['capacity'] for info in designer_info_list)
    # Count designers with any available hours.
    available_designers = len([info for info in designer_info_list if info['capacity'] > 0])
    # Assuming each designer works a 40-hour week.
    total_possible_hours = total_designers * 40
    total_assigned_hours = total_possible_hours - total_available_hours
    utilization = (total_assigned_hours / total_possible_hours * 100) if total_possible_hours else 0

    capacity_summary = (
        f"Capacity Summary: Available Designers: {available_designers} | "
        f"Available Hours: {total_available_hours:.1f} | "
        f"Utilization: {utilization:.1f}%"
    )
    
    # Compute Aggregated Project Breakdown Summary.
    total_tasks_overall = sum(sum(types.values()) for types in aggregated_breakdown.values())
    project_summary = f"Project Breakdown Summary: {total_tasks_overall} total tasks"
    
    # Compute Deadline Breakdown Summary.
    now = datetime.datetime.now()
    red = yellow = green = 0
    for d_str in aggregated_deadlines:
        try:
            d = pd.to_datetime(d_str).to_pydatetime()
        except Exception:
            continue
        delta = (d - now).days
        if delta < 0:
            continue
        elif delta < 7:
            red += 1
        elif delta < 14:
            yellow += 1
        else:
            green += 1
    deadline_summary = (
        f"Deadline Breakdown Summary: Next week: {red} | Next 2 weeks: {yellow} | Beyond 2 weeks: {green}"
    )
    
    # Create a summary table (3 rows, 1 column) with improved styling.
    summary_table = document.add_table(rows=3, cols=1)
    summary_table.style = "Table Grid"
    for i, summary in enumerate([capacity_summary, project_summary, deadline_summary]):
        cell = summary_table.rows[i].cells[0]
        cell.text = summary
        para = cell.paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in para.runs:
            run.bold = True
        set_cell_margin(cell, margin=100)
    document.add_paragraph("")  # Add a spacer paragraph
    
    # ----------------- Main Designer Capacity Table -----------------
    headers = ["Designer", "Available Hours", "Projects", "Project Type", "Deadline"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        para = hdr_cells[i].paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in para.runs:
            run.bold = True
        set_cell_shading(hdr_cells[i], fill="D9E1F2")
        set_cell_margin(hdr_cells[i], margin=100)
    set_column_widths(table, [Inches(1.5)] * len(headers))
    
    for info in designer_info_list:
        row_cells = table.add_row().cells
        row_cells[0].text = info['name']
        row_cells[1].text = f"{info['capacity']:.1f}"
        row_cells[2].text = ", ".join(sorted(info.get('projects', []))) or "None"
        row_cells[3].text = ", ".join(sorted(info.get('subtask_categories', []))) or "None"
        # Format deadline cell with red font for urgent deadlines.
        deadline_cell = row_cells[4]
        para = deadline_cell.paragraphs[0]
        para.text = ""
        deadlines = sorted(info.get('parent_deadlines', []))
        for j, d_str in enumerate(deadlines):
            run = para.add_run(d_str)
            try:
                d_dt = pd.to_datetime(d_str)
                delta = (d_dt - pd.Timestamp.now()).days
                if 0 <= delta < 7:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            except Exception:
                pass
            if j < len(deadlines) - 1:
                para.add_run(", ")
        for cell in row_cells:
            set_cell_margin(cell, margin=100)
    
    # ----------------- Aggregated Project Breakdown Section -----------------
    agg_heading = document.add_heading("Aggregated Project Breakdown", 1)
    set_collapsible(agg_heading)
    agg_table = document.add_table(rows=1, cols=3)
    agg_table.style = "Table Grid"
    agg_hdr_cells = agg_table.rows[0].cells
    agg_hdr_cells[0].text = "Project Type"
    agg_hdr_cells[1].text = "Breakdown"
    agg_hdr_cells[2].text = "Total"
    for cell in agg_hdr_cells:
        para = cell.paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in para.runs:
            run.bold = True
        set_cell_shading(cell, fill="D9E1F2")
        set_cell_margin(cell, margin=100)
    set_column_widths(agg_table, [Inches(1.5), Inches(3.5), Inches(1)])
    for project_type, projects in aggregated_breakdown.items():
        breakdown_details = ", ".join(f"{proj} ({cnt})" for proj, cnt in projects.items())
        total_count = sum(projects.values())
        row_cells = agg_table.add_row().cells
        row_cells[0].text = project_type
        row_cells[1].text = breakdown_details
        row_cells[2].text = str(total_count)
        for cell in row_cells:
            set_cell_margin(cell, margin=100)
    
    # ----------------- Deadline Breakdown Section -----------------
    deadline_heading = document.add_heading("Deadline Breakdown", 1)
    set_collapsible(deadline_heading)
    document.add_paragraph("")
    deadlines_table = document.add_table(rows=1, cols=4)
    deadlines_table.style = "Table Grid"
    dt_hdr_cells = deadlines_table.rows[0].cells
    dt_headers = ["Designer", "Project", "Project Type", "Deadline"]
    for i, header in enumerate(dt_headers):
        dt_hdr_cells[i].text = header
        para = dt_hdr_cells[i].paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in dt_hdr_cells[i].paragraphs[0].runs:
            run.bold = True
        set_cell_shading(dt_hdr_cells[i], fill="D9E1F2")
        set_cell_margin(dt_hdr_cells[i], margin=100)
    set_column_widths(deadlines_table, [Inches(1.5), Inches(2), Inches(1.5), Inches(2)])
    if not deadlines_details:
        row_cells = deadlines_table.add_row().cells
        merged = row_cells[0].merge(row_cells[3])
        merged.text = "No deadline for the week"
        for para in merged.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.bold = True
        set_cell_margin(merged, margin=100)
    else:
        for rec in deadlines_details:
            row_cells = deadlines_table.add_row().cells
            row_cells[0].text = rec.get('designer', "Unknown")
            row_cells[1].text = rec.get('project', "N/A")
            row_cells[2].text = rec.get('project_type', "N/A")
            row_cells[3].text = rec.get('deadline', "N/A")
            for cell in row_cells:
                set_cell_margin(cell, margin=100)
    
    document.add_paragraph("")
    pie_chart_buf = create_deadline_pie_chart(aggregated_deadlines)
    document.add_picture(pie_chart_buf, width=Inches(5))
    
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


def main():
    st.title("Designer Capacity & Availability Tracker")
    st.write(
        "This app retrieves data for the current work week and generates a polished Word report. "
        "The report includes a detailed table of designer capacity and availability, "
        "an aggregated project breakdown, and a deadline breakdown section with a pie chart."
    )
    
    if st.button("Run Analysis"):
        with st.spinner("Authenticating with Odoo..."):
            uid, models = authenticate_odoo()
        start_of_week, end_of_week = get_sunday_friday_range()
        start_date_str = start_of_week.strftime("%Y-%m-%d")
        end_date_str = end_of_week.strftime("%Y-%m-%d")
        with st.spinner(f"Retrieving designers for {start_date_str} to {end_date_str}..."):
            relevant_designer_ids = get_designer_ids_from_planning(models, uid, start_date_str, end_date_str)
            if not relevant_designer_ids:
                st.warning("No designers found in planning slots for this week.")
                return
        with st.spinner("Reading employee info..."):
            employees = read_employee_info(models, uid, relevant_designer_ids)
        employee_dict = {emp['id']: emp for emp in employees}
        designer_ids = list(employee_dict.keys())
        with st.spinner("Fetching timesheet data..."):
            timesheet_dict = get_all_timesheet_hours(models, uid, designer_ids, start_date_str, end_date_str)
        with st.spinner("Fetching scheduled data..."):
            scheduled_dict = get_all_scheduled_data(models, uid, designer_ids, start_date_str, end_date_str)
        with st.spinner("Fetching subtask service categories..."):
            subtask_cat_dict = get_subtask_service_categories(models, uid, designer_ids, start_date_str, end_date_str)
        with st.spinner("Fetching parent task deadlines..."):
            parent_dd_dict = get_parent_task_due_dates(models, uid, designer_ids, start_date_str, end_date_str)
        with st.spinner("Fetching project breakdown..."):
            project_breakdown_dict = get_project_breakdown(models, uid, designer_ids, start_date_str, end_date_str)
        aggregated_breakdown = aggregate_project_breakdowns(project_breakdown_dict)
        aggregated_deadlines = []
        for emp_id in designer_ids:
            aggregated_deadlines.extend(list(parent_dd_dict.get(emp_id, set())))
        deadlines_details = get_deadlines_for_week(models, uid, designer_ids, start_date_str, end_date_str)
        designer_info_list = []
        for emp_id in designer_ids:
            emp = employee_dict.get(emp_id)
            if not emp:
                continue
            name = emp.get('name', 'Unknown')
            timesheet_hours = timesheet_dict.get(emp_id, 0.0)
            sched = scheduled_dict.get(emp_id, {'hours': 0.0, 'projects': set()})
            scheduled_hours = sched['hours']
            projects = sched['projects']
            sub_cats = subtask_cat_dict.get(emp_id, set())
            parent_dds = parent_dd_dict.get(emp_id, set())
            capacity, guess = get_availability_guess_coded(name, timesheet_hours, scheduled_hours)
            designer_info_list.append({
                'name': name,
                'capacity': capacity,
                'guess': guess,
                'projects': projects,
                'subtask_categories': sub_cats,
                'parent_deadlines': parent_dds
            })
        designer_info_list.sort(key=lambda x: x['name'].lower())
        with st.spinner("Generating Word document..."):
            doc_buffer = generate_better_word_doc(designer_info_list, aggregated_breakdown, aggregated_deadlines, deadlines_details)
        st.success(f"Analysis complete for {start_date_str} to {end_date_str}!")
        st.download_button(
            label="Download Capacity Tracker",
            data=doc_buffer,
            file_name="Capacity_Tracker.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.header("Capacity Tracker Preview")
        st.write(f"**Date Range:** {start_date_str} to {end_date_str}")
        for info in designer_info_list:
            st.subheader(info['name'])
            st.write(f"**Available Hours:** {info['capacity']:.1f}, **Availability:** {info['guess']}")
            st.write(f"**Projects:** {', '.join(sorted(info['projects'])) if info['projects'] else 'None'}")
            st.write(f"**Project Type:** {', '.join(sorted(info['subtask_categories'])) if info['subtask_categories'] else 'None'}")
            st.write(f"**Deadline:** {', '.join(sorted(info['parent_deadlines'])) if info['parent_deadlines'] else 'None'}")
            st.write("---")
        st.subheader("Aggregated Project Breakdown")
        agg_total, agg_breakdown_text = format_project_breakdown_for_employee(aggregated_breakdown)
        st.write(f"Working on {agg_total} project entries: {agg_breakdown_text}")
        st.subheader("Deadline Breakdown (Pie Chart)")
        st.image(create_deadline_pie_chart(aggregated_deadlines), width=500)

if __name__ == "__main__":
    main()
